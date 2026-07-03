"""Report generation service.

Ported from old NestJS ReportService and report-generators.ts.
Generates complete .docx Word documents for house safety assessment reports.
"""

from __future__ import annotations

import uuid
from datetime import datetime
from io import BytesIO

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.exceptions import NotFoundError
from app.models.component_check import ComponentCheck
from app.models.evaluation_standard import EvaluationStandard
from app.models.report_signature import ReportSignature
from app.models.report_template import ReportTemplate
from app.models.structural_test_result import StructuralTestResult
from app.models.survey import Survey
from app.models.test_image import TestImage

CATEGORY_ORDER = ["地基基础", "上部承重结构", "围护结构", "其他"]

COMPONENT_ORDER = [
    "地基", "基础", "混凝土柱", "砖柱", "砖墙", "混凝土梁", "混凝土板", "屋架",
    "砌体自承重墙", "填充墙", "门窗洞口过梁", "挑梁", "雨棚板", "女儿墙",
    "楼地面", "屋面", "非承重墙", "门窗", "外抹灰", "内抹灰", "顶棚",
    "细木装修", "水卫", "电照", "暖通",
]


def _format_date(val: str | datetime | None) -> str:
    """Format a date as XXXX年XX月XX日."""
    if not val:
        return ""
    if isinstance(val, str):
        try:
            d = datetime.fromisoformat(val.replace("Z", "+00:00"))
        except (ValueError, TypeError):
            return val
    elif isinstance(val, datetime):
        d = val
    else:
        return str(val)
    return f"{d.year}年{d.month}月{d.day}日"


def _yn(val: bool | None) -> str:
    return "是" if val else "否"


def _arr(val: list[str] | str | None) -> str:
    if isinstance(val, list):
        return "、".join(val)
    return val or ""


def _set_cell_text(cell, text: str, bold: bool = False, size: int = 10, align=None):
    """Set cell text with formatting."""
    for paragraph in cell.paragraphs:
        paragraph.clear()
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align or WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text or "")
    run.font.size = Pt(size)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if bold:
        run.bold = True


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    """Add a centered heading."""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _add_paragraph(doc: Document, text: str, align=None, bold: bool = False,
                   spacing_before: int = 0, spacing_after: int = 0) -> None:
    """Add a formatted paragraph."""
    p = doc.add_paragraph()
    p.alignment = align or WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if bold:
        run.bold = True
    if spacing_before:
        p.paragraph_format.space_before = Pt(spacing_before)
    if spacing_after:
        p.paragraph_format.space_after = Pt(spacing_after)


async def get_full_report_data(survey_id: uuid.UUID, db: AsyncSession) -> dict:
    """Gather all report data for a survey."""
    result = await db.execute(
        select(Survey).where(Survey.id == survey_id, Survey.is_deleted == False)
    )
    survey = result.scalar_one_or_none()
    if survey is None:
        raise NotFoundError(detail="Survey not found")

    comp_result = await db.execute(
        select(ComponentCheck).where(ComponentCheck.survey_id == survey_id)
    )
    components = comp_result.scalars().all()

    test_result = await db.execute(
        select(StructuralTestResult).where(StructuralTestResult.survey_id == survey_id)
    )
    structural_tests = test_result.scalars().all()

    sig_result = await db.execute(
        select(ReportSignature).where(ReportSignature.survey_id == survey_id)
    )
    signatures = sig_result.scalars().all()

    img_result = await db.execute(
        select(TestImage)
        .where(TestImage.survey_id == survey_id)
        .order_by(TestImage.sort_order)
    )
    test_images = img_result.scalars().all()

    return {
        "survey": survey,
        "components": components,
        "structural_tests": structural_tests,
        "signatures": signatures,
        "test_images": test_images,
    }


def _sort_components(components: list[ComponentCheck]) -> list[ComponentCheck]:
    """Sort components by category order then component name order."""
    def _sort_key(c: ComponentCheck) -> tuple[int, int]:
        cat_idx = CATEGORY_ORDER.index(c.category) if c.category in CATEGORY_ORDER else 999
        comp_idx = COMPONENT_ORDER.index(c.name) if c.name in COMPONENT_ORDER else 999
        return (cat_idx, comp_idx)
    return sorted(components, key=_sort_key)


def _group_by_category(checks: list[ComponentCheck]) -> dict[str, list[ComponentCheck]]:
    """Group component checks by category."""
    grouped: dict[str, list[ComponentCheck]] = {}
    for c in checks:
        cat = c.category or "其他"
        grouped.setdefault(cat, []).append(c)
    return grouped


def _sub_group_name(checks: list[ComponentCheck]) -> dict[str, list[ComponentCheck]]:
    """Further group components by sub-name within a category."""
    sub: dict[str, list[ComponentCheck]] = {}
    for c in checks:
        name = c.name or "其他"
        if "基础" in name:
            key = "基础"
        elif "地基" in name:
            key = "地基"
        elif "柱" in name:
            key = "柱构件"
        elif "梁" in name and "过梁" not in name and "挑梁" not in name:
            key = "梁构件"
        elif "板" in name and "雨棚" not in name:
            key = "板构件"
        elif "墙" in name and "女儿" not in name and "填充" not in name and "自承重" not in name and "非承重" not in name:
            key = "墙构件"
        elif "屋架" in name or "檩条" in name:
            key = "屋架"
        elif "挑梁" in name:
            key = "挑梁"
        elif "过梁" in name:
            key = "门窗洞口过梁"
        elif "雨棚" in name:
            key = "雨棚板"
        elif "女儿墙" in name:
            key = "女儿墙"
        elif "填充墙" in name:
            key = "承担水平荷载的填充墙"
        elif "自承重墙" in name:
            key = "砌体自承重墙"
        elif "非承重墙" in name:
            key = "非承重墙"
        elif "内抹灰" in name:
            key = "内抹灰"
        elif "外抹灰" in name:
            key = "外抹灰"
        elif "地面" in name or "楼面" in name:
            key = "楼地面"
        elif "屋面" in name:
            key = "屋面"
        elif "顶棚" in name:
            key = "顶棚"
        elif "门窗" in name:
            key = "门窗"
        else:
            key = name
        sub.setdefault(key, []).append(c)
    return sub


def _collect_photos(checks: list[ComponentCheck]) -> list[dict]:
    """Collect all photos from component checks with ordering."""
    photos: list[dict] = []
    order = 1
    for c in checks:
        if c.photos:
            for url in c.photos:
                photos.append({"url": url, "order": order, "component_name": c.name})
                order += 1
    return photos


def _build_cover(doc: Document, survey: Survey) -> None:
    """Build the cover page section."""
    profile = survey.building_profile or {}
    basic = profile.get("basicInfo", {}) if isinstance(profile, dict) else {}
    bldg_name = basic.get("buildingName") or survey.address or ""

    _add_paragraph(doc, "", spacing_before=200)
    _add_paragraph(doc, "房屋安全鉴定书", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing_after=80)
    _add_paragraph(
        doc,
        f"鉴定编号：{survey.survey_no or ''}",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        spacing_before=40,
    )
    _add_paragraph(
        doc,
        f"房屋名称：{bldg_name}",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        spacing_after=120,
    )
    _add_paragraph(
        doc,
        "湖北省建筑工程质量监督检验测试中心有限公司",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        spacing_before=60,
    )
    _add_paragraph(doc, "")
    _add_paragraph(doc, "鉴定书查询方式：", spacing_before=20)
    _add_paragraph(doc, "1、查询网址：http://fgj.wuhan.gov.cn/fwagjddwjjdbgba.jhtml")
    _add_paragraph(doc, '2、关注微信公众平台"武汉住保房管"，点击下方"办事大厅"，进入"房屋安全"栏目，点击"鉴定报告查询"，可直接扫描右上方二维码查询。')
    doc.add_page_break()


def _build_overview(doc: Document, survey: Survey) -> None:
    """Build the house overview section with tables."""
    profile = survey.building_profile or {}
    basic = profile.get("basicInfo", {}) if isinstance(profile, dict) else {}
    client = profile.get("clientInfo", {}) if isinstance(profile, dict) else {}
    purpose = profile.get("purposeInfo", {}) if isinstance(profile, dict) else {}
    self_built = profile.get("selfBuiltInfo", {}) if isinstance(profile, dict) else {}
    structure = profile.get("structureInfo", {}) if isinstance(profile, dict) else {}
    history = profile.get("historyChange", {}) if isinstance(profile, dict) else {}

    report_data = survey.report_data or {}

    # ---- Table 1: Client & Building Info ----
    t1 = doc.add_table(rows=14, cols=6)
    t1.style = "Table Grid"
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER

    rows_data = [
        ["委托单位\n（委托人）", client.get("clientUnit") or client.get("clientName") or "",
         "联系人", client.get("contactPerson") or "",
         "联系电话", client.get("contactPhone") or ""],
        ["房屋名称", basic.get("buildingName") or survey.address or "",
         "产权性质", basic.get("propertyRight") or (basic.get("propertyRightUnknown") and "不详") or "",
         "产权证号", basic.get("certificateNo") or (basic.get("certificateNoUnknown") and "不详") or ""],
        ["房屋坐落", basic.get("location") or survey.address or "",
         "街道", survey.street or "",
         "社区", survey.community or ""],
        ["产权人", survey.property_owner or "",
         "建筑面积", f"{basic.get('buildArea') or survey.build_area}㎡" if (basic.get('buildArea') or survey.build_area) else "",
         "使用人", survey.property_user or ""],
        ["建成年份", basic.get("buildYearUnknown") and "不详" or (f"{basic.get('buildYear')}年" if basic.get("buildYear") else (f"{survey.build_year}年" if survey.build_year else "")),
         "地勘单位", basic.get("originalSurveyUnit") or (basic.get("originalSurveyUnitUnknown") and "不详") or "",
         "设计用途", survey.design_usage or ""],
        ["设计单位", basic.get("originalDesignUnit") or (basic.get("originalDesignUnitUnknown") and "不详") or "",
         "", "", "", ""],
        ["施工单位", basic.get("originalConstructUnit") or (basic.get("originalConstructUnitUnknown") and "不详") or "",
         "委托日期", _format_date(survey.survey_date) or "xxxx年xx月xx日",
         "监理单位", basic.get("originalSuperviseUnit") or (basic.get("originalSuperviseUnitUnknown") and "不详") or ""],
        ["查勘完成日期", _format_date(survey.inspection_date) or "xxxx年xx月xx日",
         "", "", "", ""],
        ["是否农危房改造", _yn(self_built.get("isRuralRenovation")),
         "是否优保建筑", _yn(self_built.get("isCulturalRelic")),
         "是否历史遗留办证", "否"],
        ["是否校外培训机构", _yn(self_built.get("isTrainingInstitution")),
         "是否自建房专项鉴定报告", _yn(self_built.get("isSelfBuiltReport")),
         "普查房屋编号", self_built.get("censusHouseNo") or ""],
        ["自建房排查编码", self_built.get("selfBuiltCheckCode") or "", "", "", "", ""],
        ["是否自建房", _yn(self_built.get("isSelfBuilt")),
         "是否经营性自建房", _yn(self_built.get("isCommercialSelfBuilt")),
         "", ""],
        ["现用途", _arr(purpose.get("currentPurpose")) if isinstance(purpose.get("currentPurpose"), list) else (purpose.get("currentPurpose") or ""),
         "", "", "", ""],
        ["鉴定类别", survey.survey_category or "整幢鉴定",
         "鉴定面积", f"{basic.get('buildArea') or survey.build_area}㎡" if (basic.get('buildArea') or survey.build_area) else "",
         "", ""],
    ]

    for row_idx, row_data in enumerate(rows_data):
        row = t1.rows[row_idx]
        cells = row.cells
        for ci, text in enumerate(row_data):
            if ci < len(cells):
                _set_cell_text(cells[ci], text, bold=(ci % 2 == 0), size=9)

    # Merge cells for multi-column spans
    # Row 10: 自建房排查编码 spans cols 1-5
    t1.rows[10].cells[1].merge(t1.rows[10].cells[5])
    # Row 11: 现用途 spans cols 1-5
    t1.rows[11].cells[1].merge(t1.rows[11].cells[5])

    # Add extra rows for 鉴定类别说明 and 鉴定目的
    row_cat = t1.add_row()
    _set_cell_text(row_cat.cells[0], "鉴定类别说明", bold=True, size=9)
    row_cat.cells[1].merge(row_cat.cells[5])
    _set_cell_text(row_cat.cells[1], survey.survey_category_desc or "", size=9)

    row_purpose = t1.add_row()
    _set_cell_text(row_purpose.cells[0], "鉴定目的", bold=True, size=9)
    row_purpose.cells[1].merge(row_purpose.cells[5])
    _set_cell_text(row_purpose.cells[1], survey.survey_purpose or "", size=9)

    _add_paragraph(doc, "")

    # ---- Table 2: Standards, History, Structure ----
    t2 = doc.add_table(rows=7, cols=2)
    t2.style = "Table Grid"
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER

    t2_rows = [
        ["委托人提供资料",
         report_data.get("clientProvidedMaterials") if isinstance(report_data, dict) else ""],
        ["本鉴定依据标准及规范",
         '"完标"：《房屋完损等级评定标准》（建设部城住字[84]第678号）\n'
         '"危标"：《危险房屋鉴定标准》（JGJ 125-2016）\n'
         '"规程"：《武汉市房屋安全鉴定技术规程》（DB4201/T 537-2018）'],
        ["房屋使用维修改造灾害等历史情况",
         _build_history_text(history)],
        ["房屋外部环境及周边建设施工情况",
         report_data.get("externalEnvironment") if isinstance(report_data, dict) else ""],
        ["房屋地质勘察、地基基础、主体结构及其他情况",
         f"{_build_structure_text(structure, basic)}\n\n"
         f"{report_data.get('structureCondition') if isinstance(report_data, dict) else ''}"],
        ["备注",
         "1、本鉴定书仅作为房屋安全性评估、房屋维护和危房治理的依据，不作他用；\n"
         "2、房屋需要进行维修、加固、重建的，应当到相关部门办理有关手续；\n"
         "3、本鉴定书中房屋建筑面积仅作为房屋鉴定收费依据，具体应以产权管理部门核定的建筑面积为准。"],
        ["总平面示意图", survey.site_plan_url or ""],
    ]

    for row_idx, (label, value) in enumerate(t2_rows):
        row = t2.rows[row_idx]
        _set_cell_text(row.cells[0], label, bold=True, size=9)
        # Set first column width to ~20%
        row.cells[0].width = Cm(3.5)
        _set_cell_text(row.cells[1], value, size=9)

    doc.add_page_break()


def _build_history_text(history: dict) -> str:
    """Build usage history text from history change data."""
    if not isinstance(history, dict):
        return ""
    parts = []
    map_vals = {"has": "有", "none": "无", "unknown": "不详"}
    if history.get("purposeChange"):
        parts.append(f"用途变更：{map_vals.get(history['purposeChange'], history['purposeChange'])}")
    if history.get("renovation"):
        parts.append(f"改造扩建：{map_vals.get(history['renovation'], history['renovation'])}")
    if history.get("reinforcement"):
        parts.append(f"加固修缮：{map_vals.get(history['reinforcement'], history['reinforcement'])}")
    if history.get("disaster"):
        parts.append(f"灾害：{map_vals.get(history['disaster'], history['disaster'])}")
    return "；".join(parts)


def _build_structure_text(structure: dict, basic: dict) -> str:
    """Build structure description text."""
    if not isinstance(structure, dict):
        return ""
    basic = basic if isinstance(basic, dict) else {}
    parts = []
    if structure.get("structureType"):
        parts.append(f"结构类型：{_arr(structure['structureType'])}")
    if structure.get("roofStructure"):
        parts.append(f"屋盖型式：{_arr(structure['roofStructure'])}")
    if structure.get("foundationType"):
        parts.append(f"基础型式：{_arr(structure['foundationType'])}")
    if structure.get("floorType"):
        parts.append(f"楼盖型式：{_arr(structure['floorType'])}")
    if basic.get("floorCount"):
        parts.append(f"层数：{basic['floorCount']}")
    if basic.get("totalHeight"):
        parts.append(f"檐口高度：{basic['totalHeight']}m")
    return "  ".join(parts)


def _build_component_stats(doc: Document, components: list[ComponentCheck]) -> None:
    """Build the component check statistics table."""
    counts: dict[str, int] = {cat: 0 for cat in CATEGORY_ORDER}
    for c in components:
        cat = c.category or "其他"
        counts[cat] = counts.get(cat, 0) + 1

    total = len(components)

    t = doc.add_table(rows=2, cols=6)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["构件分类", "全部"] + CATEGORY_ORDER
    for i, h in enumerate(headers):
        _set_cell_text(t.rows[0].cells[i], h, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    values = ["构件数量", str(total)] + [str(counts.get(cat, 0)) for cat in CATEGORY_ORDER]
    for i, v in enumerate(values):
        _set_cell_text(t.rows[1].cells[i], v, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    _add_paragraph(doc, "")


def _build_evaluation_table(doc: Document, components: list[ComponentCheck]) -> None:
    """Build the evaluation/damage component table."""
    grouped = _group_by_category(components)
    photo_list = _collect_photos(components)

    # Calculate total rows needed
    total_rows = 0
    for cat in CATEGORY_ORDER:
        checks = grouped.get(cat, [])
        if not checks:
            continue
        sub = _sub_group_name(checks)
        for sub_checks in sub.values():
            total_rows += len(sub_checks)

    t = doc.add_table(rows=1 + total_rows, cols=6)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ["构件名称", "", "部位", "损坏构件查勘情况", "评定结果", "评定标准条款"]
    for i, h in enumerate(headers):
        _set_cell_text(t.rows[0].cells[i], h, bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    current_row = 1
    for cat in CATEGORY_ORDER:
        checks = grouped.get(cat, [])
        if not checks:
            continue
        sub = _sub_group_name(checks)
        cat_total = sum(len(v) for v in sub.values())

        row_in_cat = 0
        for sub_name, sub_checks in sub.items():
            for check in sub_checks:
                row = t.rows[current_row]

                # Column 1: Category name (merged, only first row of category)
                if row_in_cat == 0 and cat_total > 0:
                    _set_cell_text(row.cells[0], cat, bold=True, size=8)
                    if cat_total > 1:
                        row.cells[0].merge(t.rows[current_row + cat_total - 1].cells[0])

                # Column 2: Sub-name
                _set_cell_text(row.cells[1], sub_name, size=8)

                # Column 3: Axis line / position
                _set_cell_text(row.cells[2], check.axis_line or "/", size=8)

                # Column 4: Check items description
                desc_text = ""
                if check.checked_item_ids:
                    desc_text = "；".join(str(item) for item in check.checked_item_ids)
                # Add photo references
                if check.photos:
                    photo_refs = []
                    for p in check.photos:
                        for ph in photo_list:
                            if ph["url"] == p:
                                photo_refs.append(f"见照片{ph['order']}")
                    if photo_refs:
                        desc_text += " " + "、".join(photo_refs)
                _set_cell_text(row.cells[3], desc_text, size=8)

                # Column 5: AI evaluation result
                _set_cell_text(row.cells[4], check.ai_evaluation_result or "", size=8)

                # Column 6: AI evaluation clause
                _set_cell_text(row.cells[5], check.ai_evaluation_clause or "", size=8)

                current_row += 1
                row_in_cat += 1


def _build_structural_test_table(doc: Document, structural_tests: list[StructuralTestResult],
                                 test_images: list[TestImage]) -> None:
    """Build the structural test results table."""
    test = structural_tests[0] if structural_tests else None

    rows_needed = 7  # base rows
    if test_images:
        rows_needed += len(test_images)

    t = doc.add_table(rows=rows_needed, cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    base_rows = [
        ["检测单位", test.test_unit if test else "湖北省建筑工程质量监督检验测试中心有限公司"],
        ["检测资质证书号", test.certificate_no if test else "鄂建检字第OJ06001号"],
        ["检测人员", test.test_personnel if test else ""],
        ["检测报告编号", test.report_no if test else ""],
        ["主要检测内容", test.main_test_content if test else "1、结构现状检查；2、房屋整体倾斜观测。"],
        ["检测依据标准", test.test_standards if test else (
            "《危险房屋鉴定标准》JGJ 125-2016；《武汉市房屋安全鉴定技术规程》DB4201/T 537-2018；"
            "《建筑结构检测技术标准》GB/T 50344-2019；《建筑变形测量规范》JGJ 8-2016；"
            "《房屋完损等级评定标准》（建设部城住字[84]第678号）。"
        )],
        ["主要检测成果", test.test_results_summary if test else ""],
    ]

    for i, (label, value) in enumerate(base_rows):
        _set_cell_text(t.rows[i].cells[0], label, bold=True, size=9)
        t.rows[i].cells[0].width = Cm(4)
        _set_cell_text(t.rows[i].cells[1], value, size=9)

    # Add test image rows
    if test_images:
        img_row = len(base_rows)
        for img in test_images:
            label_map = {
                "elevation_front": "正立面图",
                "elevation_back": "背立面图",
                "floor_plan": "平面图",
                "site_plan": "总平面图",
            }
            label = img.label or label_map.get(img.type, "图片")
            _set_cell_text(t.rows[img_row].cells[0], label, bold=True, size=9)
            _set_cell_text(t.rows[img_row].cells[1], f"[图片: {img.image_url}]", size=9,
                           align=WD_ALIGN_PARAGRAPH.CENTER)
            img_row += 1


def _build_conclusion_table(doc: Document, survey: Survey,
                            structural_tests: list[StructuralTestResult],
                            signatures: list[ReportSignature]) -> None:
    """Build the conclusion and signature table."""
    ai = survey.ai_reasoning_result or {}
    if not isinstance(ai, dict):
        ai = {}

    test = structural_tests[0] if structural_tests else None

    sig_map: dict[str, ReportSignature] = {}
    for s in signatures:
        if s.type:
            sig_map[s.type] = s

    t = doc.add_table(rows=8, cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    conclusion_rows = [
        ["房屋损坏情况综述", ai.get("basicEvaluation") or test.damage_summary if test else ""],
        ["房屋损坏原因分析", test.cause_analysis if test else ""],
        ["鉴定结论及处理意见",
         f"{ai.get('conclusion') or test.conclusion if test else ''}\n\n"
         f"建议：{ai.get('suggestion') or test.handling_suggestion if test else ''}"],
        ["安全等级", ai.get("riskLevel") or (f"{test.safety_level}级" if test and test.safety_level else "")],
        ["鉴定人", sig_map.get("appraiser") and (sig_map["appraiser"].signatory_name or "[签字]") or ""],
        ["审核人", sig_map.get("reviewer") and (sig_map["reviewer"].signatory_name or "[签字]") or ""],
        ["签发人", sig_map.get("issuer") and (sig_map["issuer"].signatory_name or "[签字]") or ""],
        ["单位公章", sig_map.get("seal") and "[盖章]" or ""],
    ]

    for i, (label, value) in enumerate(conclusion_rows):
        _set_cell_text(t.rows[i].cells[0], label, bold=True, size=9)
        t.rows[i].cells[0].width = Cm(4)
        _set_cell_text(t.rows[i].cells[1], value, size=9)


def _build_photo_section(doc: Document, components: list[ComponentCheck]) -> None:
    """Build the inspection photos section."""
    photos = _collect_photos(components)
    if not photos:
        return

    _add_heading(doc, "第五部分：鉴定房屋照片", level=1)
    _add_paragraph(doc, "")

    for i in range(0, len(photos), 2):
        t = doc.add_table(rows=2, cols=2)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER

        for col in range(2):
            idx = i + col
            if idx < len(photos):
                photo = photos[idx]
                _set_cell_text(t.rows[0].cells[col], f"照片 {photo['order']}", size=9,
                               align=WD_ALIGN_PARAGRAPH.CENTER)
                _set_cell_text(t.rows[1].cells[col], f"[图片: {photo['url']}]", size=9,
                               align=WD_ALIGN_PARAGRAPH.CENTER)

        _add_paragraph(doc, "")


async def generate_report_docx(survey_id: uuid.UUID, db: AsyncSession) -> BytesIO:
    """Generate a complete .docx report for a survey.

    Returns a BytesIO stream containing the .docx file bytes.
    """
    data = await get_full_report_data(survey_id, db)

    survey: Survey = data["survey"]
    components: list[ComponentCheck] = data["components"]
    structural_tests: list[StructuralTestResult] = data["structural_tests"]
    signatures: list[ReportSignature] = data["signatures"]
    test_images: list[TestImage] = data["test_images"]

    sorted_components = _sort_components(components)

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # ---- Cover Page ----
    _build_cover(doc, survey)

    # ---- Part 1: House Overview ----
    _add_heading(doc, "第一部分：房屋概况", level=1)
    _build_overview(doc, survey)

    # ---- Part 2: Damage Component Evaluation ----
    _add_heading(doc, "第二部分：地基及损坏结构构件评定表", level=1)
    _build_component_stats(doc, sorted_components)
    _build_evaluation_table(doc, sorted_components)
    _add_paragraph(doc, "")
    doc.add_page_break()

    # ---- Part 3: Structural Test Results ----
    _add_heading(doc, "第三部分：结构检测成果表", level=1)
    _build_structural_test_table(doc, structural_tests, test_images)
    _add_paragraph(doc, "")
    doc.add_page_break()

    # ---- Part 4: Conclusion ----
    _add_heading(doc, "第四部分：房屋鉴定结论和处理建议", level=1)
    _build_conclusion_table(doc, survey, structural_tests, signatures)
    _add_paragraph(doc, "")

    # ---- Part 5: Photos ----
    _build_photo_section(doc, sorted_components)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
