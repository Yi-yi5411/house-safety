"""Original record generation service.

Ported from old NestJS OriginalRecordService.
Generates .docx documents for original inspection records (原始记录).
"""

from __future__ import annotations

import uuid
from datetime import datetime
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.exceptions import NotFoundError
from app.models.component_check import ComponentCheck
from app.models.survey import Survey
from app.models.structural_test_result import StructuralTestResult
from app.models.report_signature import ReportSignature


def _set_cell_text(cell, text: str, bold: bool = False, size: int = 10, align=None):
    """Set cell text with formatting."""
    for paragraph in cell.paragraphs:
        paragraph.clear()
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align or WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(str(text) if text else "")
    run.font.size = Pt(size)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if bold:
        run.bold = True


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    """Add a centered heading."""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _add_paragraph(doc: Document, text: str, align=None, bold: bool = False,
                   spacing_before: int = 0, spacing_after: int = 0) -> None:
    """Add a formatted paragraph."""
    p = doc.add_paragraph()
    p.alignment = align or WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text or "")
    run.font.size = Pt(12)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if bold:
        run.bold = True
    if spacing_before:
        p.paragraph_format.space_before = Pt(spacing_before)
    if spacing_after:
        p.paragraph_format.space_after = Pt(spacing_after)


CATEGORY_ORDER = ["地基基础", "上部承重结构", "围护结构", "其他"]


async def get_original_record_data(survey_id: uuid.UUID, db: AsyncSession) -> dict:
    """Gather original record data for a survey."""
    result = await db.execute(
        select(Survey).where(Survey.id == survey_id, Survey.is_deleted == False)
    )
    survey = result.scalar_one_or_none()
    if survey is None:
        raise NotFoundError(detail="Survey not found")

    comp_result = await db.execute(
        select(ComponentCheck).where(ComponentCheck.survey_id == survey_id)
    )
    components = comp_result.scalars().all()

    # Also get structural test results and signatures for completeness
    test_result = await db.execute(
        select(StructuralTestResult).where(StructuralTestResult.survey_id == survey_id)
    )
    tests = test_result.scalars().all()

    sig_result = await db.execute(
        select(ReportSignature).where(ReportSignature.survey_id == survey_id)
    )
    signatures = sig_result.scalars().all()

    return {
        "survey": survey,
        "components": components,
        "structural_tests": tests,
        "signatures": signatures,
    }


def _build_building_info_table(doc: Document, survey: Survey) -> None:
    """Build the building basic info table."""
    t = doc.add_table(rows=10, cols=4)
    t.style = "Table Grid"

    info_rows = [
        ["房屋地址", survey.address or "", "建造年份", f"{survey.build_year}年" if survey.build_year else ""],
        ["结构类型", survey.structure_type or "", "层数", str(survey.floor_count) if survey.floor_count else ""],
        ["建筑面积", f"{survey.build_area}㎡" if survey.build_area else "", "檐口高度", survey.eaves_height or ""],
        ["产权人", survey.property_owner or "", "使用人", survey.property_user or ""],
        ["委托人", survey.client_name or "", "联系电话", survey.contact_phone or ""],
        ["设计用途", survey.design_usage or "", "现用途", survey.current_usage or ""],
        ["鉴定类别", survey.survey_category or "整幢鉴定", "鉴定编号", survey.survey_no or ""],
        ["鉴定目的", survey.survey_purpose or "", "委托日期", str(survey.entrust_date) if survey.entrust_date else ""],
        ["查勘日期", str(survey.inspection_date) if survey.inspection_date else "", "完成日期", str(survey.inspection_complete_date) if survey.inspection_complete_date else ""],
        ["街道", survey.street or "", "社区", survey.community or ""],
    ]

    for row_idx, row_data in enumerate(info_rows):
        for col_idx, text in enumerate(row_data):
            bold = col_idx % 2 == 0
            _set_cell_text(t.rows[row_idx].cells[col_idx], text, bold=bold, size=9)


def _build_component_check_table(doc: Document, components: list[ComponentCheck]) -> None:
    """Build the component inspection record table."""
    if not components:
        _add_paragraph(doc, "无构件检查记录")
        return

    t = doc.add_table(rows=1 + len(components), cols=7)
    t.style = "Table Grid"

    headers = ["序号", "构件名称", "类别", "部位(轴线)", "损坏描述", "评定结果", "照片数量"]
    for i, h in enumerate(headers):
        _set_cell_text(t.rows[0].cells[i], h, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    for idx, comp in enumerate(components):
        row = t.rows[idx + 1]
        _set_cell_text(row.cells[0], str(idx + 1), size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(row.cells[1], comp.name or "", size=9)
        _set_cell_text(row.cells[2], comp.category or "", size=9)
        _set_cell_text(row.cells[3], comp.axis_line or "/", size=9)
        _set_cell_text(row.cells[4], comp.damage_description or "", size=9)
        _set_cell_text(row.cells[5], comp.ai_evaluation_result or "", size=9)
        _set_cell_text(row.cells[6], str(len(comp.photos)) if comp.photos else "0",
                       size=9, align=WD_ALIGN_PARAGRAPH.CENTER)


def _build_structural_test_section(doc: Document, tests: list[StructuralTestResult]) -> None:
    """Build structural test results section."""
    if not tests:
        return

    _add_heading(doc, "结构检测结果", level=2)

    for test in tests:
        t = doc.add_table(rows=6, cols=2)
        t.style = "Table Grid"

        test_rows = [
            ["检测单位", test.test_unit or ""],
            ["检测资质证书号", test.certificate_no or ""],
            ["检测人员", test.test_personnel or ""],
            ["检测报告编号", test.report_no or ""],
            ["主要检测内容", test.main_test_content or ""],
            ["检测结果综述", test.test_results_summary or ""],
        ]

        for ri, (label, value) in enumerate(test_rows):
            _set_cell_text(t.rows[ri].cells[0], label, bold=True, size=9)
            t.rows[ri].cells[0].width = Cm(3.5)
            _set_cell_text(t.rows[ri].cells[1], value, size=9)

        # Add extra rows for the 5 new fields if present
        extra_fields = [
            ("损伤综述", test.damage_summary),
            ("原因分析", test.cause_analysis),
            ("结论", test.conclusion),
            ("处理建议", test.handling_suggestion),
            ("安全等级", f"{test.safety_level}级" if test.safety_level else None),
        ]
        for label, value in extra_fields:
            if value:
                row = t.add_row()
                _set_cell_text(row.cells[0], label, bold=True, size=9)
                _set_cell_text(row.cells[1], value, size=9)

        _add_paragraph(doc, "")


def _build_signature_section(doc: Document, signatures: list[ReportSignature]) -> None:
    """Build signature section."""
    if not signatures:
        return

    _add_heading(doc, "签名信息", level=2)

    t = doc.add_table(rows=len(signatures), cols=4)
    t.style = "Table Grid"

    headers = ["类型", "签名人", "签名图片", "签名日期"]
    for i, h in enumerate(headers):
        _set_cell_text(t.rows[0].cells[i], h, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    type_names = {"appraiser": "鉴定人", "reviewer": "审核人", "issuer": "签发人", "seal": "公章"}
    for idx, sig in enumerate(signatures):
        _set_cell_text(t.rows[idx].cells[0], type_names.get(sig.type, sig.type or ""), size=9)
        _set_cell_text(t.rows[idx].cells[1], sig.signatory_name or "", size=9)
        _set_cell_text(t.rows[idx].cells[2], "已上传" if sig.image_url else "无", size=9,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(t.rows[idx].cells[3], str(sig.sign_date) if sig.sign_date else "", size=9)


async def generate_original_record_docx(survey_id: uuid.UUID, db: AsyncSession) -> BytesIO:
    """Generate a complete .docx original record for a survey.

    Returns a BytesIO stream containing the .docx file bytes.
    """
    data = await get_original_record_data(survey_id, db)

    survey: Survey = data["survey"]
    components: list[ComponentCheck] = data["components"]
    tests: list[StructuralTestResult] = data["structural_tests"]
    signatures: list[ReportSignature] = data["signatures"]

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Title
    _add_heading(doc, "房屋安全鉴定原始记录", level=0)
    _add_paragraph(doc, f"记录编号：{survey.survey_no or ''}", align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "")

    # ---- Part 1: Building Basic Info ----
    _add_heading(doc, "一、房屋基本信息", level=1)
    _build_building_info_table(doc, survey)
    _add_paragraph(doc, "")

    # AI reasoning result info
    ai = survey.ai_reasoning_result or {}
    if isinstance(ai, dict) and ai:
        _add_heading(doc, "二、AI推理结果", level=1)
        _add_paragraph(doc, f"鉴定结论: {ai.get('conclusion', '')}")
        _add_paragraph(doc, f"基础评定: {ai.get('basicEvaluation', '')}")
        _add_paragraph(doc, f"风险等级: {ai.get('riskLevel', '')}")
        _add_paragraph(doc, f"建议: {ai.get('suggestion', '')}")
        _add_paragraph(doc, "")

    # ---- Part 2: Component Check Records ----
    heading_num = "三" if (isinstance(ai, dict) and ai) else "二"
    _add_heading(doc, f"{heading_num}、构件检查记录", level=1)
    _build_component_check_table(doc, components)
    _add_paragraph(doc, "")

    # ---- Part 3: Structural Test Results ----
    if tests:
        _build_structural_test_section(doc, tests)

    # ---- Part 4: Signatures ----
    if signatures:
        _build_signature_section(doc, signatures)

    # ---- Footer ----
    _add_paragraph(doc, "")
    _add_paragraph(doc, f"鉴定单位：湖北省建筑工程质量监督检验测试中心有限公司")
    _add_paragraph(doc, f"记录日期：{datetime.now().strftime('%Y年%m月%d日')}")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
